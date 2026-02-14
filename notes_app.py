import streamlit as st
import re
import docx
import PyPDF2
import os
import google.generativeai as genai
from google.api_core.exceptions import ResourceExhausted

# Configure API key
genai.configure(api_key=os.getenv("GOOGLE_API_KEY"))

# App title
st.title("📘 AI Notes Assistant")

# Tabs
tab1, tab2, tab3 = st.tabs([
    "✍️ Generate Notes from Topic",
    "📝 Summarize My Notes",
    "❓ Generate Study Questions"
])

# ---------------- TAB 1 ----------------
with tab1:
    topic = st.text_input("Enter a topic for notes:")

    if st.button("Generate Notes", key="gen_notes_btn"):
        if topic.strip():
            with st.spinner("Generating notes..."):
                model = genai.GenerativeModel("gemini-2.0-flash")
                try:
                    response = model.generate_content(
                        f"Write detailed, structured notes on {topic}."
                    )
                    notes_text = response.text

                    st.subheader("📖 Generated Notes")
                    st.write(notes_text)

                    st.download_button(
                        label="💾 Download Notes",
                        data=notes_text,
                        file_name="notes.txt",
                        mime="text/plain"
                    )

                except ResourceExhausted:
                    st.error("""
🚫 API quota is over.

You have exceeded the daily request limit.
Please try again later or upgrade your API plan.
""")

                except Exception as e:
                    st.error(f"⚠️ Unexpected error: {str(e)}")
        else:
            st.warning("⚠️ Please enter a topic first.")

# ---------------- TAB 2 ----------------
with tab2:
    st.subheader("📄 Summarize Your Notes")

    uploaded_file = st.file_uploader(
        "Upload your notes (.pdf or .doc/.docx)",
        type=["pdf", "doc", "docx"],
        key="upload_summary"
    )

    notes_to_summarize = ""

    if uploaded_file is not None:
        if uploaded_file.name.endswith(".pdf"):
            reader = PyPDF2.PdfReader(uploaded_file)
            for page in reader.pages:
                notes_to_summarize += page.extract_text() + "\n"
        elif uploaded_file.name.endswith((".doc", ".docx")):
            doc = docx.Document(uploaded_file)
            for para in doc.paragraphs:
                notes_to_summarize += para.text + "\n"
    else:
        notes_to_summarize = st.text_area(
            "Or paste your notes here:",
            height=200
        )

    if st.button("Summarize Notes", key="summarize_btn"):
        if notes_to_summarize.strip():
            with st.spinner("Summarizing your notes..."):
                model = genai.GenerativeModel("gemini-2.0-flash")
                try:
                    response = model.generate_content(
                        f"Summarize the following notes:\n\n{notes_to_summarize}"
                    )
                    summary_text = response.text

                    st.subheader("📝 Summary")
                    st.write(summary_text)

                    st.download_button(
                        label="💾 Download Summary",
                        data=summary_text,
                        file_name="summary.txt",
                        mime="text/plain"
                    )

                except ResourceExhausted:
                    st.error("""
🚫 API quota is over.

You have exceeded the daily request limit.
Please try again later or upgrade your API plan.
""")

                except Exception as e:
                    st.error(f"⚠️ Unexpected error: {str(e)}")
        else:
            st.warning("⚠️ Please upload a file or paste some notes.")

# ---------------- TAB 3 ----------------
with tab3:
    uploaded_notes = st.file_uploader(
        "Upload notes for MCQs (.pdf or .doc/.docx)",
        type=["pdf", "doc", "docx"],
        key="upload_mcqs"
    )

    notes_for_questions = ""

    if uploaded_notes is not None:
        if uploaded_notes.name.endswith(".pdf"):
            reader = PyPDF2.PdfReader(uploaded_notes)
            for page in reader.pages:
                notes_for_questions += page.extract_text() + "\n"
        elif uploaded_notes.name.endswith((".doc", ".docx")):
            doc = docx.Document(uploaded_notes)
            for para in doc.paragraphs:
                notes_for_questions += para.text + "\n"
    else:
        notes_for_questions = st.text_area(
            "Or paste your notes here to generate MCQs:",
            height=200
        )

    if "mcqs" not in st.session_state:
        st.session_state.mcqs = []
        st.session_state.user_answers = []
        st.session_state.show_results = False

    if st.button("Generate MCQs", key="gen_mcqs_btn"):
        if notes_for_questions.strip():
            num_words = len(notes_for_questions.split())
            num_mcqs = max(10, min(20, num_words // 80))

            with st.spinner(f"Generating {num_mcqs} multiple-choice questions..."):
                model = genai.GenerativeModel("gemini-2.0-flash")
                try:
                    response = model.generate_content(
                        f"From these notes, generate {num_mcqs} multiple-choice questions (MCQs). "
                        "Each question should have 4 options (A, B, C, D). "
                        "Clearly mark the correct answer with 'Answer: X'. "
                        "Format exactly like:\n\n"
                        "Q1. Question text?\nA) ...\nB) ...\nC) ...\nD) ...\nAnswer: B\n\n"
                        f"Notes:\n{notes_for_questions}"
                    )

                    raw_mcqs = response.text
                    questions = re.split(r"Q\d+\.", raw_mcqs)
                    mcqs_list = []

                    for q in questions[1:]:
                        parts = q.strip().split("\n")
                        if len(parts) < 6:
                            continue
                        question_text = parts[0]
                        options = [p for p in parts[1:5]]
                        answer_line = [p for p in parts if p.startswith("Answer:")]
                        correct_answer = (
                            answer_line[0].split(":")[-1].strip()
                            if answer_line else None
                        )

                        mcqs_list.append({
                            "question": question_text,
                            "options": options,
                            "answer": correct_answer
                        })

                    st.session_state.mcqs = mcqs_list
                    st.session_state.user_answers = [""] * len(mcqs_list)
                    st.session_state.show_results = False

                except ResourceExhausted:
                    st.error("""
🚫 API quota is over.

You have exceeded the daily request limit.
Please try again later or upgrade your API plan.
""")

                except Exception as e:
                    st.error(f"⚠️ Unexpected error: {str(e)}")
        else:
            st.warning("⚠️ Please upload or paste your notes first.")

    if st.session_state.mcqs:
        st.subheader("🎯 Quiz Time!")

        for idx, mcq in enumerate(st.session_state.mcqs):
            st.session_state.user_answers[idx] = st.radio(
                f"Q{idx+1}. {mcq['question']}",
                options=mcq["options"],
                key=f"q_{idx}"
            )

        if st.button("Submit Quiz"):
            st.session_state.show_results = True

        if st.session_state.show_results:
            st.subheader("📊 Quiz Results")
            score = 0

            for i, mcq in enumerate(st.session_state.mcqs):
                user_ans = st.session_state.user_answers[i]
                correct_ans = mcq["answer"]

                st.markdown(f"**Q{i+1}: {mcq['question']}**")

                if user_ans and correct_ans and user_ans.startswith(correct_ans):
                    st.success(f"✅ Correct! Your answer: {user_ans}")
                    score += 1
                else:
                    st.error(f"❌ Wrong. Your answer: {user_ans} | Correct answer: {correct_ans}")

                st.write("---")

            st.subheader(f"🏆 Your Score: {score}/{len(st.session_state.mcqs)}")






